#!/usr/bin/env python3
"""
Post-process a pandoc-generated DOCX to match JMIS/EJIS LaTeX formatting.

Matches jmisstyle.sty specifications:
  - Times New Roman 12pt body, Arial 10pt tables/figures
  - Double-spaced body text (2.0), single-spaced headings
  - 1-inch margins on US Letter
  - 0.5in first-line paragraph indent
  - Heading 1: centered, bold, uppercase, 12pt, black
  - Heading 2: left-aligned, bold, 12pt, black
  - Fully justified body text
  - All Word theme references stripped

Usage:
    python scripts/format_docx.py input.docx [output.docx]

If output is omitted, overwrites the input file.

Requirements:
    pip install python-docx lxml
"""

import sys
import re
import copy
import zipfile
import tempfile
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

# ---------------------------------------------------------------------------
# Constants — derived from jmisstyle.sty
# ---------------------------------------------------------------------------
BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 12            # \fontsize{12}{14}
TABLE_FONT = "Arial"
TABLE_SIZE_PT = 10           # \fontsize{10}{12}
LINE_SPACING_DOUBLE = 480    # double-space in twips (240 twips per line × 2)
LINE_SPACING_SINGLE = 240   # single-space in twips
HEADING_SPACING_BEFORE = 120 # 6pt before headings (only gap between sections)
HEADING_SPACING_AFTER = 0    # 0pt after headings
MARGIN_INCHES = 1.0
INDENT_INCHES = 0.5
HANGING_INDENT_INCHES = 0.25
BLACK = "000000"

THEME_FONT_ATTRS = [
    qn("w:asciiTheme"), qn("w:hAnsiTheme"),
    qn("w:cstheme"), qn("w:eastAsiaTheme"),
]
THEME_COLOR_ATTRS = [
    qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade"),
]


def strip_theme_refs(element):
    """Remove every theme-font and theme-color attribute under *element*."""
    for rfonts in element.iter(qn("w:rFonts")):
        for attr in THEME_FONT_ATTRS:
            rfonts.attrib.pop(attr, None)
    for color in element.iter(qn("w:color")):
        for attr in THEME_COLOR_ATTRS:
            color.attrib.pop(attr, None)


def set_font_on_rfonts(rfonts, name):
    """Set all four font slots to *name*."""
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def make_clean_rpr(font=BODY_FONT, size_pt=BODY_SIZE_PT, bold=False,
                   italic=False, color=BLACK):
    """Build a <w:rPr> element with explicit values and no theme refs."""
    rpr = etree.SubElement(etree.Element("tmp"), qn("w:rPr"))
    rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    set_font_on_rfonts(rfonts, font)
    if bold:
        etree.SubElement(rpr, qn("w:b"))
        etree.SubElement(rpr, qn("w:bCs"))
    i = etree.SubElement(rpr, qn("w:i"))
    i.set(qn("w:val"), "1" if italic else "0")
    iCs = etree.SubElement(rpr, qn("w:iCs"))
    iCs.set(qn("w:val"), "1" if italic else "0")
    c = etree.SubElement(rpr, qn("w:color"))
    c.set(qn("w:val"), color)
    sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), str(size_pt * 2))  # half-points
    szCs = etree.SubElement(rpr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(size_pt * 2))
    return rpr


def replace_rpr(run_elem, new_rpr):
    """Swap the <w:rPr> child of a <w:r> element."""
    old = run_elem.find(qn("w:rPr"))
    if old is not None:
        run_elem.remove(old)
    run_elem.insert(0, copy.deepcopy(new_rpr))


# ---------------------------------------------------------------------------
# Style-level fixes
# ---------------------------------------------------------------------------

def fix_style(style_elem, font, size_pt, bold, italic, color, alignment=None):
    """Rewrite a style definition's rPr and optionally pPr alignment."""
    # Rebuild rPr
    old_rpr = style_elem.find(qn("w:rPr"))
    if old_rpr is not None:
        style_elem.remove(old_rpr)
    new_rpr = make_clean_rpr(font, size_pt, bold, italic, color)
    style_elem.append(new_rpr)

    # Alignment
    if alignment is not None:
        ppr = style_elem.find(qn("w:pPr"))
        if ppr is None:
            ppr = etree.SubElement(style_elem, qn("w:pPr"))
        jc = ppr.find(qn("w:jc"))
        if jc is None:
            jc = etree.SubElement(ppr, qn("w:jc"))
        jc.set(qn("w:val"), alignment)


def fix_styles(doc):
    """Fix all relevant built-in style definitions."""
    style_configs = {
        "Title":          dict(font=BODY_FONT, size_pt=16, bold=True,
                               italic=False, color=BLACK, alignment="center"),
        "Title Char":     dict(font=BODY_FONT, size_pt=16, bold=True,
                               italic=False, color=BLACK),
        "TitleChar":      dict(font=BODY_FONT, size_pt=16, bold=True,
                               italic=False, color=BLACK),
        "Abstract Title": dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=False, color=BLACK, alignment="center"),
        "Abstract":       dict(font=BODY_FONT, size_pt=12, bold=False,
                               italic=True, color=BLACK, alignment="both"),
        "Heading 1":      dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=False, color=BLACK, alignment="center"),
        "Heading 1 Char": dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=False, color=BLACK),
        "Heading 2":      dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=False, color=BLACK, alignment="left"),
        "Heading 2 Char": dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=False, color=BLACK),
        "Heading 3":      dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=True, color=BLACK, alignment="left"),
        "Heading 3 Char": dict(font=BODY_FONT, size_pt=12, bold=True,
                               italic=True, color=BLACK),
    }
    for s in doc.styles:
        cfg = style_configs.get(s.name)
        if cfg:
            fix_style(s.element, **cfg)
            print(f"  Fixed style: {s.name}")

    # Fix Bibliography style — controls BIBLIOGRAPHY field output spacing
    for s in doc.styles:
        if s.name == "Bibliography":
            fix_style(s.element, font=BODY_FONT, size_pt=BODY_SIZE_PT,
                      bold=False, italic=False, color=BLACK, alignment="left")
            ppr = s.element.find(qn("w:pPr"))
            if ppr is None:
                ppr = etree.SubElement(s.element, qn("w:pPr"))
            hanging = int(0.25 * 1440)
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=120)
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = etree.SubElement(ppr, qn("w:ind"))
            ind.set(qn("w:left"), str(hanging))
            ind.set(qn("w:hanging"), str(hanging))
            print(f"  Fixed style: {s.name}")

    # Strip theme refs from every remaining style
    for s in doc.styles:
        strip_theme_refs(s.element)


# ---------------------------------------------------------------------------
# Theme-level fixes
# ---------------------------------------------------------------------------

def fix_theme(doc):
    """Set the theme's major/minor latin fonts to Times New Roman."""
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    for rel in doc.part.rels.values():
        if "theme" in rel.reltype:
            tree = etree.fromstring(rel.target_part.blob)
            for latin in tree.iter(f"{{{ns_a}}}latin"):
                latin.set("typeface", BODY_FONT)
            for ea in tree.iter(f"{{{ns_a}}}ea"):
                ea.set("typeface", "")
            for cs in tree.iter(f"{{{ns_a}}}cs"):
                cs.set("typeface", "")
            rel.target_part._blob = etree.tostring(
                tree, xml_declaration=True, encoding="UTF-8", standalone=True
            )
            print("  Fixed theme fonts")
            break


# ---------------------------------------------------------------------------
# Section (page layout) fixes
# ---------------------------------------------------------------------------

def fix_page_layout(doc):
    """Set 1-inch margins on US Letter for every section."""
    margin_twips = int(MARGIN_INCHES * 1440)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES)
        section.right_margin = Inches(MARGIN_INCHES)
    print("  Fixed page layout: US Letter, 1-inch margins")


# ---------------------------------------------------------------------------
# Paragraph-level fixes
# ---------------------------------------------------------------------------

def set_spacing(ppr, line=None, before=None, after=None, line_rule="auto"):
    """Set <w:spacing> on a pPr element."""
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(ppr, qn("w:spacing"))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), line_rule)
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))


def set_indent(ppr, first_line_twips=None, hanging_twips=None, left_twips=None):
    """Set indent on a pPr element. Supports first-line and hanging."""
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = etree.SubElement(ppr, qn("w:ind"))
    if first_line_twips is not None:
        ind.attrib.pop(qn("w:hanging"), None)
        ind.set(qn("w:firstLine"), str(first_line_twips))
    if hanging_twips is not None:
        ind.attrib.pop(qn("w:firstLine"), None)
        ind.set(qn("w:hanging"), str(hanging_twips))
        if left_twips is None:
            left_twips = hanging_twips
    if left_twips is not None:
        ind.set(qn("w:left"), str(left_twips))


def ensure_ppr(para):
    """Get or create <w:pPr> on a Paragraph object."""
    elem = para._element
    ppr = elem.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.SubElement(elem, qn("w:pPr"))
        elem.remove(ppr)
        elem.insert(0, ppr)
    return ppr


def is_in_table(para):
    """Check if paragraph element is inside a table cell."""
    parent = para._element.getparent()
    while parent is not None:
        if parent.tag == qn("w:tc"):
            return True
        parent = parent.getparent()
    return False


def set_alignment(ppr, val):
    """Set paragraph alignment (center, left, both, etc.)."""
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = etree.SubElement(ppr, qn("w:jc"))
    jc.set(qn("w:val"), val)


def clear_first_indent(ppr):
    """Remove first-line indent, set to 0."""
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = etree.SubElement(ppr, qn("w:ind"))
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLine"), "0")


def apply_font_to_runs(para, font=BODY_FONT, size_pt=BODY_SIZE_PT):
    """Set font and size on all runs, stripping theme refs."""
    for run in para.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = etree.SubElement(run._element, qn("w:rPr"))
            run._element.insert(0, rpr)
        strip_theme_refs(rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, qn("w:rFonts"))
        set_font_on_rfonts(rfonts, font)
        sz = rpr.find(qn("w:sz"))
        if sz is None:
            sz = etree.SubElement(rpr, qn("w:sz"))
        sz.set(qn("w:val"), str(size_pt * 2))
        szCs = rpr.find(qn("w:szCs"))
        if szCs is None:
            szCs = etree.SubElement(rpr, qn("w:szCs"))
        szCs.set(qn("w:val"), str(size_pt * 2))
        color_el = rpr.find(qn("w:color"))
        if color_el is not None:
            for attr in THEME_COLOR_ATTRS:
                color_el.attrib.pop(attr, None)
            color_el.set(qn("w:val"), BLACK)


def fix_keywords_and_pagebreak(doc, tex_path=None):
    """Inject keywords from .tex source (pandoc drops them from flushleft/titlepage)
    and add a page break before the first Heading 1 (Introduction)."""
    keywords_text = None
    if tex_path and Path(tex_path).exists():
        with open(tex_path, encoding='utf-8') as f:
            tex = f.read()
        m = re.search(
            r'\\textbf\{Keywords:\}\s*(.+?)(?:\n\\end\{flushleft\}|\n\n)',
            tex, re.DOTALL
        )
        if m:
            keywords_text = "Keywords: " + m.group(1).strip()

    paras = doc.paragraphs
    first_h1_idx = None
    abstract_end_idx = None

    for i, para in enumerate(paras):
        style = para.style.name if para.style else ''
        if style == 'Abstract':
            abstract_end_idx = i
        if style == 'Heading 1' and first_h1_idx is None:
            first_h1_idx = i

    if keywords_text and abstract_end_idx is not None:
        target_idx = abstract_end_idx + 1
        if target_idx < len(paras):
            target = paras[target_idx]
            if not target.text.strip():
                target.text = ""
                from docx.oxml import OxmlElement
                bold_rpr = make_clean_rpr(BODY_FONT, BODY_SIZE_PT, bold=True,
                                          italic=False, color=BLACK)
                normal_rpr = make_clean_rpr(BODY_FONT, BODY_SIZE_PT, bold=False,
                                            italic=False, color=BLACK)
                r_bold = OxmlElement("w:r")
                r_bold.append(copy.deepcopy(bold_rpr))
                t_bold = OxmlElement("w:t")
                t_bold.text = "Keywords:"
                t_bold.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                r_bold.append(t_bold)
                target._element.append(r_bold)

                rest = keywords_text[len("Keywords:"):]
                r_normal = OxmlElement("w:r")
                r_normal.append(copy.deepcopy(normal_rpr))
                t_normal = OxmlElement("w:t")
                t_normal.text = rest
                t_normal.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                r_normal.append(t_normal)
                target._element.append(r_normal)

                ppr = ensure_ppr(target)
                clear_first_indent(ppr)
                set_spacing(ppr, line=LINE_SPACING_DOUBLE, before=0, after=0)
                set_alignment(ppr, "both")
                print(f"  Injected keywords paragraph")

    if first_h1_idx is not None:
        h1_para = paras[first_h1_idx]
        ppr = ensure_ppr(h1_para)
        pb = ppr.find(qn("w:pageBreakBefore"))
        if pb is None:
            etree.SubElement(ppr, qn("w:pageBreakBefore"))
        print(f"  Added page break before first heading (Introduction)")


def fix_paragraphs(doc, anonymize=False):
    """Apply formatting to every paragraph based on its style."""
    indent_twips = int(INDENT_INCHES * 1440)  # 720 twips = 0.5 inch
    hanging_twips = int(HANGING_INDENT_INCHES * 1440)  # 360 twips = 0.25 inch
    counts = {"h1": 0, "h2": 0, "title": 0, "author": 0,
              "abstract": 0, "bib": 0, "body": 0}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        ppr = ensure_ppr(para)
        in_table = is_in_table(para)

        # --- Title (16pt, bold, centered, ALL CAPS) ---
        if style_name == "Title":
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=0)
            set_alignment(ppr, "center")
            clear_first_indent(ppr)
            rpr_t = make_clean_rpr(BODY_FONT, 16, bold=True, italic=False,
                                   color=BLACK)
            for run in para.runs:
                replace_rpr(run._element, rpr_t)
                if run.text:
                    run.text = run.text.upper()
            counts["title"] += 1

        # --- Author: remove if anonymizing, else style as a centered block ---
        elif style_name == "Author":
            if anonymize:
                para._element.getparent().remove(para._element)
            else:
                set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=120)
                set_alignment(ppr, "center")
                clear_first_indent(ppr)
                rpr_a = make_clean_rpr(BODY_FONT, BODY_SIZE_PT, bold=False,
                                       italic=False, color=BLACK)
                for run in para.runs:
                    replace_rpr(run._element, rpr_a)
            counts["author"] += 1

        # --- Abstract Title ---
        elif style_name == "Abstract Title":
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=120, after=0)
            set_alignment(ppr, "center")
            clear_first_indent(ppr)
            rpr_t = make_clean_rpr(BODY_FONT, 12, bold=True, italic=False,
                                   color=BLACK)
            for run in para.runs:
                replace_rpr(run._element, rpr_t)
            counts["abstract"] += 1

        # --- Abstract body ---
        elif style_name == "Abstract":
            set_spacing(ppr, line=LINE_SPACING_DOUBLE, before=0, after=0)
            set_alignment(ppr, "both")
            clear_first_indent(ppr)
            rpr_t = make_clean_rpr(BODY_FONT, 12, bold=False, italic=True,
                                   color=BLACK)
            for run in para.runs:
                replace_rpr(run._element, rpr_t)
            counts["abstract"] += 1

        # --- Heading 1 ---
        elif style_name == "Heading 1":
            set_spacing(ppr, line=LINE_SPACING_SINGLE,
                        before=HEADING_SPACING_BEFORE,
                        after=HEADING_SPACING_AFTER)
            set_alignment(ppr, "center")
            clear_first_indent(ppr)
            rpr_t = make_clean_rpr(BODY_FONT, 12, bold=True, italic=False,
                                   color=BLACK)
            for run in para.runs:
                replace_rpr(run._element, rpr_t)
                if run.text:
                    run.text = run.text.upper()

            # REFERENCES: page break before, unnumbered
            import re
            clean_text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "",
                                para.text.strip().upper())
            if clean_text in UNNUMBERED_HEADINGS:
                pb = ppr.find(qn("w:pageBreakBefore"))
                if pb is None:
                    etree.SubElement(ppr, qn("w:pageBreakBefore"))

            counts["h1"] += 1

        # --- Heading 2 ---
        elif style_name == "Heading 2":
            set_spacing(ppr, line=LINE_SPACING_SINGLE,
                        before=HEADING_SPACING_BEFORE,
                        after=int(HEADING_SPACING_AFTER * 0.5))
            set_alignment(ppr, "left")
            clear_first_indent(ppr)
            rpr_t = make_clean_rpr(BODY_FONT, 12, bold=True, italic=False,
                                   color=BLACK)
            for run in para.runs:
                replace_rpr(run._element, rpr_t)
            counts["h2"] += 1

        # --- Heading 3 ---
        elif style_name.startswith("Heading 3"):
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=240, after=60)
            clear_first_indent(ppr)
            rpr_t = make_clean_rpr(BODY_FONT, 12, bold=True, italic=True,
                                   color=BLACK)
            for run in para.runs:
                replace_rpr(run._element, rpr_t)

        # --- Bibliography / References ---
        elif style_name == "Bibliography":
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=120)
            set_alignment(ppr, "left")
            set_indent(ppr, hanging_twips=hanging_twips,
                       left_twips=hanging_twips)
            apply_font_to_runs(para)
            counts["bib"] += 1

        # --- Body text (everything else) ---
        else:
            if not in_table:
                set_spacing(ppr, line=LINE_SPACING_DOUBLE, before=0, after=0)
                set_alignment(ppr, "both")

                is_list = ppr.find(qn("w:numPr")) is not None
                is_special = style_name in (
                    "Caption", "Block Text", "List Paragraph",
                    "Figure", "Table", "Subtitle",
                    "Image Caption", "Table Caption", "Captioned Figure",
                )
                is_keywords = para.text.startswith("Keywords:")
                if not is_list and not is_special and not is_keywords:
                    set_indent(ppr, first_line_twips=indent_twips)
                if is_keywords:
                    clear_first_indent(ppr)

            apply_font_to_runs(para)
            counts["body"] += 1

    parts = [f'{counts["h1"]} H1', f'{counts["h2"]} H2',
             f'{counts["title"]} title', f'{counts["author"]} author',
             f'{counts["abstract"]} abstract', f'{counts["bib"]} bib',
             f'{counts["body"]} body']
    print(f"  Fixed paragraphs: {', '.join(parts)}")


# ---------------------------------------------------------------------------
# Heading numbering (1, 1.1, 1.1.1)
# ---------------------------------------------------------------------------

UNNUMBERED_HEADINGS = {"REFERENCES", "ACKNOWLEDGMENTS", "APPENDIX"}


def _is_unnumbered_heading(para):
    """Check if heading should be excluded from numbering."""
    text = para.text.strip().upper()
    # Strip any existing number prefix for comparison
    import re
    text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)
    return text in UNNUMBERED_HEADINGS


def _create_numbering_part(doc):
    """Create or get the numbering part with a multilevel list definition
    linked to Heading 1/2/3 styles."""
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    num_part = doc.part.numbering_part
    numbering = num_part.element

    # Check if we already have our abstract num (id=99)
    for an in numbering.findall(qn("w:abstractNum")):
        if an.get(qn("w:abstractNumId")) == "99":
            return 99  # Already created

    # Create abstract numbering definition
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "99")
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    # Level 0: "1." for Heading 1
    lvl0 = OxmlElement("w:lvl")
    lvl0.set(qn("w:ilvl"), "0")
    start0 = OxmlElement("w:start")
    start0.set(qn("w:val"), "1")
    lvl0.append(start0)
    fmt0 = OxmlElement("w:numFmt")
    fmt0.set(qn("w:val"), "decimal")
    lvl0.append(fmt0)
    txt0 = OxmlElement("w:lvlText")
    txt0.set(qn("w:val"), "%1.")
    lvl0.append(txt0)
    jc0 = OxmlElement("w:lvlJc")
    jc0.set(qn("w:val"), "left")
    lvl0.append(jc0)
    pstyle0 = OxmlElement("w:pStyle")
    pstyle0.set(qn("w:val"), "Heading1")
    lvl0.append(pstyle0)
    ppr0 = OxmlElement("w:pPr")
    ind0 = OxmlElement("w:ind")
    ind0.set(qn("w:left"), "0")
    ind0.set(qn("w:firstLine"), "0")
    ppr0.append(ind0)
    lvl0.append(ppr0)
    suff0 = OxmlElement("w:suff")
    suff0.set(qn("w:val"), "space")
    lvl0.append(suff0)
    abstract.append(lvl0)

    # Level 1: "1.1" for Heading 2
    lvl1 = OxmlElement("w:lvl")
    lvl1.set(qn("w:ilvl"), "1")
    start1 = OxmlElement("w:start")
    start1.set(qn("w:val"), "1")
    lvl1.append(start1)
    fmt1 = OxmlElement("w:numFmt")
    fmt1.set(qn("w:val"), "decimal")
    lvl1.append(fmt1)
    txt1 = OxmlElement("w:lvlText")
    txt1.set(qn("w:val"), "%1.%2")
    lvl1.append(txt1)
    jc1 = OxmlElement("w:lvlJc")
    jc1.set(qn("w:val"), "left")
    lvl1.append(jc1)
    pstyle1 = OxmlElement("w:pStyle")
    pstyle1.set(qn("w:val"), "Heading2")
    lvl1.append(pstyle1)
    ppr1 = OxmlElement("w:pPr")
    ind1 = OxmlElement("w:ind")
    ind1.set(qn("w:left"), "0")
    ind1.set(qn("w:firstLine"), "0")
    ppr1.append(ind1)
    lvl1.append(ppr1)
    suff1 = OxmlElement("w:suff")
    suff1.set(qn("w:val"), "space")
    lvl1.append(suff1)
    abstract.append(lvl1)

    # Level 2: "1.1.1" for Heading 3
    lvl2 = OxmlElement("w:lvl")
    lvl2.set(qn("w:ilvl"), "2")
    start2 = OxmlElement("w:start")
    start2.set(qn("w:val"), "1")
    lvl2.append(start2)
    fmt2 = OxmlElement("w:numFmt")
    fmt2.set(qn("w:val"), "decimal")
    lvl2.append(fmt2)
    txt2 = OxmlElement("w:lvlText")
    txt2.set(qn("w:val"), "%1.%2.%3")
    lvl2.append(txt2)
    jc2 = OxmlElement("w:lvlJc")
    jc2.set(qn("w:val"), "left")
    lvl2.append(jc2)
    pstyle2 = OxmlElement("w:pStyle")
    pstyle2.set(qn("w:val"), "Heading3")
    lvl2.append(pstyle2)
    ppr2 = OxmlElement("w:pPr")
    ind2 = OxmlElement("w:ind")
    ind2.set(qn("w:left"), "0")
    ind2.set(qn("w:firstLine"), "0")
    ppr2.append(ind2)
    lvl2.append(ppr2)
    suff2 = OxmlElement("w:suff")
    suff2.set(qn("w:val"), "space")
    lvl2.append(suff2)
    abstract.append(lvl2)

    # Insert abstract num before any w:num elements
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        numbering.insert(list(numbering).index(first_num), abstract)
    else:
        numbering.append(abstract)

    # Create concrete numbering reference
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "99")
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), "99")
    num.append(abs_ref)
    numbering.append(num)

    return 99


def fix_heading_numbering(doc):
    """Apply Word's built-in multilevel numbering to Heading 1/2/3.

    Uses a numbering definition linked to heading styles so numbers
    auto-update when headings are reordered. Skips unnumbered headings
    like REFERENCES.
    """
    import re
    from docx.oxml import OxmlElement

    num_id = _create_numbering_part(doc)
    numbered = 0

    for para in doc.paragraphs:
        if not para.style or para.style.name not in (
            "Heading 1", "Heading 2", "Heading 3"
        ):
            continue

        # Strip pandoc-injected number prefix (may span multiple runs:
        # run0="1", run1="\t", run2="Introduction")
        runs = para.runs
        if runs:
            while runs and re.fullmatch(r'\d+(?:\.\d+)*\.?', runs[0].text.strip()):
                runs[0]._element.getparent().remove(runs[0]._element)
                runs = para.runs
            if runs and runs[0].text.strip() == '':
                runs[0]._element.getparent().remove(runs[0]._element)
                runs = para.runs
            if runs:
                runs[0].text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "",
                                       runs[0].text)

        ppr = ensure_ppr(para)

        if _is_unnumbered_heading(para):
            # Remove any numPr for unnumbered headings
            existing = ppr.find(qn("w:numPr"))
            if existing is not None:
                ppr.remove(existing)
            continue

        # Apply numbering: set numPr with numId and ilvl
        level = int(para.style.name[-1]) - 1
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            ppr.remove(num_pr)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_pr.append(ilvl)
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_el)
        # Insert numPr right after pStyle if present
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is not None:
            pstyle.addnext(num_pr)
        else:
            ppr.insert(0, num_pr)
        numbered += 1

    print(f"  Numbered {numbered} headings (field codes, numId={num_id})")


# ---------------------------------------------------------------------------
# Table fixes
# ---------------------------------------------------------------------------

def _format_caption_para(para, prefix):
    """Apply shared caption formatting: set Arial 10pt, single-spaced,
    no indent, bold label, and keepNext/keepLines to glue to figure/table."""
    has_seq = para._element.find(f'.//{qn("w:fldChar")}') is not None
    if not has_seq:
        if para.runs and not para.runs[0].text.startswith(prefix.split(".")[0]):
            para.runs[0].text = prefix + para.runs[0].text
        elif not para.runs:
            from docx.oxml import OxmlElement
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = prefix
            r.append(t)
            para._element.append(r)

    # Paragraph formatting: single-spaced, no indent, left-aligned
    ppr = ensure_ppr(para)
    set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=120)
    clear_first_indent(ppr)
    set_alignment(ppr, "left")

    # Keep with next paragraph (glue caption to figure/table)
    if ppr.find(qn("w:keepNext")) is None:
        etree.SubElement(ppr, qn("w:keepNext"))
    if ppr.find(qn("w:keepLines")) is None:
        etree.SubElement(ppr, qn("w:keepLines"))

    # Set all runs to Arial 10pt, bold the label prefix
    caption_rpr = make_clean_rpr(TABLE_FONT, TABLE_SIZE_PT, bold=True,
                                  italic=False, color=BLACK)
    body_rpr = make_clean_rpr(TABLE_FONT, TABLE_SIZE_PT, bold=False,
                               italic=False, color=BLACK)
    for i, run in enumerate(para.runs):
        if i == 0:
            replace_rpr(run._element, caption_rpr)
        else:
            replace_rpr(run._element, body_rpr)


def fix_captions(doc):
    """Add 'Figure N.' / 'Table N.' prefixes and format caption paragraphs.

    Captions: Arial 10pt, single-spaced, no indent, bold label,
    keepNext to glue to the adjacent figure/table.
    """
    fig_num = 0
    tab_num = 0
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""

        # Format Captioned Figure paragraphs (the figure container)
        if style_name == "Captioned Figure":
            ppr = ensure_ppr(para)
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=120, after=0)
            clear_first_indent(ppr)
            set_alignment(ppr, "center")
            if ppr.find(qn("w:keepNext")) is None:
                etree.SubElement(ppr, qn("w:keepNext"))

        elif style_name == "Image Caption":
            fig_num += 1
            _format_caption_para(para, f"Figure {fig_num}. ")

        elif style_name == "Table Caption":
            tab_num += 1
            _format_caption_para(para, f"Table {tab_num}. ")

    print(f"  Fixed captions: {fig_num} figures, {tab_num} tables")


def replace_image(doc, image_path, rel_id=None):
    """Replace or insert the figure image in the document.

    If an existing image relationship is found, replaces its blob.
    Otherwise, finds the 'Captioned Figure' paragraph and inserts the image there.
    """
    from docx.shared import Inches as InchesShared
    img_data = Path(image_path).read_bytes()

    # Try to replace existing image
    if rel_id is None:
        for rid, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                rel_id = rid
                break

    if rel_id and rel_id in doc.part.rels:
        rel = doc.part.rels[rel_id]
        rel.target_part._blob = img_data
        print(f"  Replaced image {rel_id} with {image_path} ({len(img_data)} bytes)")
        return

    # No existing image — insert a new paragraph with the image before the
    # first Image Caption paragraph (pandoc can't render TikZ, so no figure exists)
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        if style_name == 'Image Caption':
            # Use doc.add_picture to create a paragraph with the image at the end,
            # then move it before the caption paragraph
            from docx.shared import Inches as _Inches
            pic_shape = doc.add_picture(image_path, width=_Inches(6.0))
            # add_picture returns an InlineShape; navigate up: inline→drawing→r→p
            pic_p_elem = pic_shape._inline.getparent().getparent().getparent()

            # Format the figure paragraph
            ppr = ensure_ppr_elem(pic_p_elem)
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=120, after=0)
            clear_first_indent(ppr)
            set_alignment(ppr, 'center')
            etree.SubElement(ppr, qn('w:keepNext'))

            # Move from end of document to before caption
            body = doc.element.find(qn('w:body'))
            body.remove(pic_p_elem)
            caption_elem = para._element
            body.insert(list(body).index(caption_elem), pic_p_elem)

            print(f"  Inserted figure image before caption ({len(img_data)} bytes)")
            return

    print(f"  Warning: could not find image relationship or Image Caption paragraph")


def ensure_ppr_elem(p_elem):
    """Ensure a <w:pPr> child exists on a raw <w:p> element."""
    ppr = p_elem.find(qn('w:pPr'))
    if ppr is None:
        ppr = etree.SubElement(p_elem, qn('w:pPr'))
        p_elem.insert(0, ppr)
    return ppr


TABLE_CELL_MARGIN_TWIPS = int(0.02 * 1440)  # 0.02 inches = ~29 twips


def _make_border(style="single", sz="4", color="000000"):
    """Build a border element dict for reuse."""
    return {"style": style, "sz": sz, "color": color}


def _set_cell_borders(tc, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell element."""
    tcp = tc.find(qn("w:tcPr"))
    if tcp is None:
        tcp = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcp)
    borders = tcp.find(qn("w:tcBorders"))
    if borders is not None:
        tcp.remove(borders)
    borders = etree.SubElement(tcp, qn("w:tcBorders"))
    for side, cfg in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if cfg:
            el = etree.SubElement(borders, qn(f"w:{side}"))
            el.set(qn("w:val"), cfg["style"])
            el.set(qn("w:sz"), cfg["sz"])
            el.set(qn("w:color"), cfg["color"])
            el.set(qn("w:space"), "0")


def fix_tables(doc):
    """Apply Table Style 2 look: horizontal rules, bold header with thick
    bottom border, left-aligned, 0.02in cell margins, Arial 10pt, single-spaced.
    Also removes empty placeholder tables (e.g. pandoc's TikZ stub)."""
    count = 0
    removed = 0
    thin = _make_border("single", "4", "000000")
    thick = _make_border("single", "12", "000000")
    none = _make_border("nil", "0", "000000")

    for table in list(doc.tables):
        # Remove empty tables (all cells contain only whitespace)
        all_text = ''.join(
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
        )
        if not all_text:
            tbl = table._tbl
            tbl.getparent().remove(tbl)
            removed += 1
            continue
        tbl = table._tbl
        tpr = tbl.find(qn("w:tblPr"))
        if tpr is None:
            tpr = etree.SubElement(tbl, qn("w:tblPr"))
            tbl.insert(0, tpr)

        # Remove any existing table style (we apply formatting directly)
        tbl_style = tpr.find(qn("w:tblStyle"))
        if tbl_style is not None:
            tpr.remove(tbl_style)

        # Left-align the table
        jc = tpr.find(qn("w:jc"))
        if jc is None:
            jc = etree.SubElement(tpr, qn("w:jc"))
        jc.set(qn("w:val"), "start")

        # Remove table-level borders (we set per-cell)
        tbl_borders = tpr.find(qn("w:tblBorders"))
        if tbl_borders is not None:
            tpr.remove(tbl_borders)

        # Set cell margins to 0.02 inches on all sides
        cell_mar = tpr.find(qn("w:tblCellMar"))
        if cell_mar is not None:
            tpr.remove(cell_mar)
        cell_mar = etree.SubElement(tpr, qn("w:tblCellMar"))
        for side in ("top", "left", "bottom", "right"):
            el = etree.SubElement(cell_mar, qn(f"w:{side}"))
            el.set(qn("w:w"), str(TABLE_CELL_MARGIN_TWIPS))
            el.set(qn("w:type"), "dxa")

        # Ensure tblLook enables header row and banded rows
        tbl_look = tpr.find(qn("w:tblLook"))
        if tbl_look is None:
            tbl_look = etree.SubElement(tpr, qn("w:tblLook"))
        tbl_look.set(qn("w:val"), "04A0")
        tbl_look.set(qn("w:firstRow"), "1")
        tbl_look.set(qn("w:lastRow"), "0")
        tbl_look.set(qn("w:firstColumn"), "0")
        tbl_look.set(qn("w:lastColumn"), "0")
        tbl_look.set(qn("w:noHBand"), "0")
        tbl_look.set(qn("w:noVBand"), "1")

        rows = list(table.rows)
        for row_idx, row in enumerate(rows):
            is_header = (row_idx == 0)
            is_last = (row_idx == len(rows) - 1)

            for cell in row.cells:
                tc = cell._tc
                # Table Style 2: top rule on header, thick bottom on header,
                # thin bottom on last row, horizontal rules between rows,
                # no vertical borders
                if is_header:
                    _set_cell_borders(tc,
                                      top=thin,
                                      bottom=thick,
                                      left=none,
                                      right=none)
                elif is_last:
                    _set_cell_borders(tc,
                                      top=thin,
                                      bottom=thin,
                                      left=none,
                                      right=none)
                else:
                    _set_cell_borders(tc,
                                      top=thin,
                                      bottom=none,
                                      left=none,
                                      right=none)

                # Set Arial 10pt, single-spaced, left-aligned on all cell content
                for para in cell.paragraphs:
                    ppr = para._element.find(qn("w:pPr"))
                    if ppr is None:
                        ppr = etree.SubElement(para._element, qn("w:pPr"))
                        para._element.insert(0, ppr)
                    set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=0)
                    set_alignment(ppr, "left")
                    clear_first_indent(ppr)
                    for run in para.runs:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is None:
                            rpr = etree.SubElement(run._element, qn("w:rPr"))
                            run._element.insert(0, rpr)
                        strip_theme_refs(rpr)
                        rfonts = rpr.find(qn("w:rFonts"))
                        if rfonts is None:
                            rfonts = etree.SubElement(rpr, qn("w:rFonts"))
                        set_font_on_rfonts(rfonts, TABLE_FONT)
                        sz = rpr.find(qn("w:sz"))
                        if sz is None:
                            sz = etree.SubElement(rpr, qn("w:sz"))
                        sz.set(qn("w:val"), str(TABLE_SIZE_PT * 2))
                        szCs = rpr.find(qn("w:szCs"))
                        if szCs is None:
                            szCs = etree.SubElement(rpr, qn("w:szCs"))
                        szCs.set(qn("w:val"), str(TABLE_SIZE_PT * 2))
        count += 1
    msg = f"  Fixed {count} tables (Table Style 2 borders, left-aligned, single-spaced, Arial {TABLE_SIZE_PT}pt)"
    if removed:
        msg += f", removed {removed} empty tables"
    print(msg)


# ---------------------------------------------------------------------------
# Default paragraph style (Normal)
# ---------------------------------------------------------------------------

def fix_normal_style(doc):
    """Set the Normal style to Times New Roman 12pt, double-spaced, justified."""
    for s in doc.styles:
        if s.name == "Normal":
            strip_theme_refs(s.element)
            rpr = s.element.find(qn("w:rPr"))
            if rpr is None:
                rpr = etree.SubElement(s.element, qn("w:rPr"))
            # Font
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = etree.SubElement(rpr, qn("w:rFonts"))
            for attr in THEME_FONT_ATTRS:
                rfonts.attrib.pop(attr, None)
            set_font_on_rfonts(rfonts, BODY_FONT)
            # Size
            sz = rpr.find(qn("w:sz"))
            if sz is None:
                sz = etree.SubElement(rpr, qn("w:sz"))
            sz.set(qn("w:val"), str(BODY_SIZE_PT * 2))
            szCs = rpr.find(qn("w:szCs"))
            if szCs is None:
                szCs = etree.SubElement(rpr, qn("w:szCs"))
            szCs.set(qn("w:val"), str(BODY_SIZE_PT * 2))

            # Spacing
            ppr = s.element.find(qn("w:pPr"))
            if ppr is None:
                ppr = etree.SubElement(s.element, qn("w:pPr"))
            set_spacing(ppr, line=LINE_SPACING_DOUBLE, before=0, after=0)

            # Justified
            jc = ppr.find(qn("w:jc"))
            if jc is None:
                jc = etree.SubElement(ppr, qn("w:jc"))
            jc.set(qn("w:val"), "both")

            print("  Fixed Normal style")
            break


# ---------------------------------------------------------------------------
# BibTeX parsing and LaTeX-to-Unicode conversion
# ---------------------------------------------------------------------------

LATEX_TO_UNICODE = {
    r'\"u': 'ü', r'\"a': 'ä', r'\"o': 'ö',
    r'\"U': 'Ü', r'\"A': 'Ä', r'\"O': 'Ö',
    r'\`e': 'è', r'\`a': 'à', r"\'e": 'é',
    r"\'a": 'á', r'\AA': 'Å', r'\aa': 'å',
    r'\~n': 'ñ', r'\c{c}': 'ç',
}


def _latex_to_unicode(text):
    """Convert LaTeX special characters to Unicode."""
    for latex, uni in LATEX_TO_UNICODE.items():
        text = text.replace('{' + latex + '}', uni)
        text = text.replace(latex, uni)
    text = re.sub(r'\{([^}])\}', r'\1', text)
    text = text.replace('{', '').replace('}', '')
    text = text.replace('``', '“').replace("''", '”')
    text = text.replace('--', '–')
    return text


def parse_bib(bib_path):
    """Parse a .bib file into structured entries."""
    with open(bib_path, encoding='utf-8') as f:
        content = f.read()

    entries = {}
    for m in re.finditer(
        r'@(\w+)\{(\w+),\s*(.*?)\n\}', content, re.DOTALL
    ):
        entry_type, key, body = m.group(1).lower(), m.group(2), m.group(3)
        fields = {}
        for fm in re.finditer(
            r'(\w+)\s*=\s*\{((?:[^{}]|\{[^{}]*\})*)\}', body
        ):
            fields[fm.group(1).lower()] = fm.group(2).strip()

        fields['_type'] = entry_type
        fields['_key'] = key
        entries[key] = fields

    return entries


def _parse_authors(author_str):
    """Parse BibTeX author string into list of (last, first) tuples.

    Handles corporate authors wrapped in double braces: {{Name}} → (Name, '').
    Must check for corporate braces BEFORE calling _latex_to_unicode (which
    strips all braces).
    """
    raw = author_str.strip()
    # Handle corporate authors: {{Name}} or {Name} — entire name is last name
    if raw.startswith('{') and raw.endswith('}'):
        corp_name = _latex_to_unicode(raw.strip('{}').strip())
        return [(corp_name, '')]

    parts = re.split(r'\s+and\s+', raw)
    authors = []
    for part in parts:
        part = part.strip()
        # Check for braced corporate author within multi-author field
        if part.startswith('{') and part.endswith('}'):
            authors.append((_latex_to_unicode(part.strip('{}').strip()), ''))
        elif ',' in part:
            last, first = part.split(',', 1)
            authors.append((_latex_to_unicode(last.strip()),
                            _latex_to_unicode(first.strip())))
        else:
            part = _latex_to_unicode(part)
            tokens = part.split()
            if len(tokens) >= 2:
                authors.append((' '.join(tokens[:-1]), tokens[-1]))
            else:
                authors.append((part, ''))
    return authors


NAME_PARTICLES = {'von', 'van', 'de', 'del', 'der', 'den', 'vom', 'zum', 'zur',
                   'di', 'da', 'dos', 'das', 'la', 'le', 'du', 'des'}


def _cite_surname(last_name):
    """Get the citation-form surname, stripping lowercase name particles.

    CSL treats 'von', 'de', 'vom', etc. as non-dropping particles that
    are omitted in short (in-text) citation forms. E.g. 'vom Brocke' → 'Brocke'.
    """
    tokens = last_name.split()
    # Strip leading lowercase particles
    while len(tokens) > 1 and tokens[0].lower() in NAME_PARTICLES:
        tokens.pop(0)
    return ' '.join(tokens)


def _rendered_citation(bib_entry, maxnames=3):
    """Predict how pandoc/citeproc renders an in-text citation for a single key.

    Pandoc citeproc shows all names for ≤3 authors; for 4+ uses 'et al.'.
    maxnames controls the threshold (default 3 matches citeproc default).
    """
    author_str = bib_entry.get('author', '')
    year = bib_entry.get('year', 'n.d.')
    authors = _parse_authors(author_str)

    def surname(a):
        return _cite_surname(a[0])

    if len(authors) == 0:
        # Institutional/corporate author
        title = bib_entry.get('title', '')
        if title:
            return f"{_latex_to_unicode(title)} {year}"
        return f"({year})"
    elif len(authors) == 1:
        return f"{surname(authors[0])} {year}"
    elif len(authors) == 2:
        return f"{surname(authors[0])} and {surname(authors[1])} {year}"
    else:
        if len(authors) <= maxnames:
            names = ', '.join(surname(a) for a in authors[:-1])
            return f"{names}, and {surname(authors[-1])} {year}"
        else:
            return f"{surname(authors[0])} et al. {year}"


def _build_citation_map(bib_entries, tex_path):
    """Build mapping from rendered citation group text to list of bib keys.

    Uses the .tex source to know which keys are grouped together in each
    \\parencite{} command, then predicts the rendered form from .bib data.
    Returns list of (rendered_text_inside_parens, [keys]) tuples.
    """
    with open(tex_path, encoding='utf-8') as f:
        tex = f.read()

    cite_commands = re.findall(
        r'\\parencite\{([^}]+)\}', tex
    )

    groups = []
    for cmd in cite_commands:
        keys = [k.strip() for k in cmd.split(',')]
        parts = []
        for key in keys:
            if key in bib_entries:
                parts.append(_rendered_citation(bib_entries[key]))
            else:
                parts.append(key)
        rendered = '; '.join(parts)
        groups.append((rendered, keys))

    return groups


# ---------------------------------------------------------------------------
# Word bibliography source XML (custom XML part)
# ---------------------------------------------------------------------------

BIB_NS = "http://schemas.openxmlformats.org/officeDocument/2006/bibliography"

_BIB_TYPE_MAP = {
    'article': 'JournalArticle',
    'book': 'Book',
    'inproceedings': 'ConferenceProceedings',
    'incollection': 'BookSection',
    'phdthesis': 'Report',
    'mastersthesis': 'Report',
    'techreport': 'Report',
    'online': 'InternetSite',
    'misc': 'Misc',
}


def _bib_to_sources_xml(bib_entries):
    """Convert parsed bib entries to Word bibliography sources XML."""
    ns = BIB_NS
    root = etree.Element(f'{{{ns}}}Sources',
                         nsmap={'b': ns})
    root.set('SelectedStyle', r'\APASixthEditionOfficeOnline.xsl')
    root.set('StyleName', 'APA')
    root.set('Version', '6')

    for key, entry in bib_entries.items():
        src = etree.SubElement(root, f'{{{ns}}}Source')
        etree.SubElement(src, f'{{{ns}}}Tag').text = key
        source_type = _BIB_TYPE_MAP.get(entry['_type'], 'Misc')
        etree.SubElement(src, f'{{{ns}}}SourceType').text = source_type

        title = _latex_to_unicode(entry.get('title', ''))
        etree.SubElement(src, f'{{{ns}}}Title').text = title

        year = entry.get('year', '')
        if year:
            etree.SubElement(src, f'{{{ns}}}Year').text = year

        authors_raw = entry.get('author', '')
        if authors_raw:
            author_el = etree.SubElement(src, f'{{{ns}}}Author')
            author_list = etree.SubElement(author_el, f'{{{ns}}}Author')
            name_list = etree.SubElement(author_list, f'{{{ns}}}NameList')
            for last, first in _parse_authors(authors_raw):
                person = etree.SubElement(name_list, f'{{{ns}}}Person')
                etree.SubElement(person, f'{{{ns}}}Last').text = last
                if first:
                    etree.SubElement(person, f'{{{ns}}}First').text = first

        journal = entry.get('journaltitle', '') or entry.get('journal', '')
        if journal:
            etree.SubElement(src, f'{{{ns}}}JournalName').text = _latex_to_unicode(journal)

        publisher = entry.get('publisher', '')
        if publisher:
            etree.SubElement(src, f'{{{ns}}}Publisher').text = _latex_to_unicode(publisher)

        location = entry.get('location', '') or entry.get('address', '')
        if location:
            etree.SubElement(src, f'{{{ns}}}City').text = _latex_to_unicode(location)

        volume = entry.get('volume', '')
        if volume:
            etree.SubElement(src, f'{{{ns}}}Volume').text = volume

        issue = entry.get('number', '')
        if issue:
            etree.SubElement(src, f'{{{ns}}}Issue').text = issue

        pages = entry.get('pages', '').replace('--', '-')
        if pages:
            etree.SubElement(src, f'{{{ns}}}Pages').text = pages

        doi = entry.get('doi', '')
        if doi:
            etree.SubElement(src, f'{{{ns}}}DOI').text = doi

        url = entry.get('url', '')
        if url:
            etree.SubElement(src, f'{{{ns}}}URL').text = url

        edition = entry.get('edition', '')
        if edition:
            etree.SubElement(src, f'{{{ns}}}Edition').text = edition

    return root


def _inject_bibliography_xml(docx_path, bib_entries):
    """Inject Word bibliography sources as custom XML into the DOCX zip."""
    sources_xml = _bib_to_sources_xml(bib_entries)
    sources_bytes = etree.tostring(sources_xml, xml_declaration=True,
                                    encoding='UTF-8', standalone=True,
                                    pretty_print=True)

    item_props = (
        '<?xml version="1.0" encoding="UTF-8" standalone="no"?>\r\n'
        '<ds:datastoreItem ds:itemID="{B1D504B5-31AA-4F5A-8102-1B37B8A2B400}"'
        ' xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml">'
        '<ds:schemaRefs>'
        '<ds:schemaRef ds:uri="http://schemas.openxmlformats.org/officeDocument/2006/bibliography"/>'
        '</ds:schemaRefs></ds:datastoreItem>'
    ).encode('utf-8')

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()

    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp.name, 'w', zipfile.ZIP_DEFLATED) as zout:

        existing = set(zin.namelist())
        ct_xml = None

        doc_rels_path = 'word/_rels/document.xml.rels'
        for item in zin.namelist():
            data = zin.read(item)
            if item == '[Content_Types].xml':
                ct_xml = data
                continue
            if item == doc_rels_path:
                # Will be rewritten with new relationship below
                continue
            zout.writestr(item, data)

        item_num = 1
        while f'customXml/item{item_num}.xml' in existing:
            item_num += 1

        zout.writestr(f'customXml/item{item_num}.xml', sources_bytes)
        zout.writestr(f'customXml/itemProps{item_num}.xml', item_props)

        if ct_xml:
            ct_tree = etree.fromstring(ct_xml)
            ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
            needed = [
                (f'/customXml/item{item_num}.xml', 'application/xml'),
                (f'/customXml/itemProps{item_num}.xml',
                 'application/vnd.openxmlformats-officedocument.customXmlProperties+xml'),
            ]
            for pn, ct in needed:
                exists = any(
                    el.get('PartName') == pn
                    for el in ct_tree.findall(f'{{{ct_ns}}}Override')
                )
                if not exists:
                    etree.SubElement(ct_tree, f'{{{ct_ns}}}Override',
                                    PartName=pn, ContentType=ct)
            zout.writestr('[Content_Types].xml',
                          etree.tostring(ct_tree, xml_declaration=True,
                                         encoding='UTF-8', standalone=True))

        if doc_rels_path in existing:
            rels_data = zin.read(doc_rels_path)
        else:
            rels_data = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
            ).encode('utf-8')

        rels_tree = etree.fromstring(rels_data)
        rels_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
        max_id = 0
        for rel in rels_tree:
            rid = rel.get('Id', '')
            m = re.match(r'rId(\d+)', rid)
            if m:
                max_id = max(max_id, int(m.group(1)))

        new_rid = f'rId{max_id + 1}'
        etree.SubElement(rels_tree, f'{{{rels_ns}}}Relationship',
                         Id=new_rid,
                         Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml',
                         Target=f'../customXml/item{item_num}.xml')

        zout.writestr(doc_rels_path,
                      etree.tostring(rels_tree, xml_declaration=True,
                                     encoding='UTF-8', standalone=True))

    shutil.move(tmp.name, docx_path)
    print(f"  Injected {len(bib_entries)} bibliography sources into Reference Manager")


# ---------------------------------------------------------------------------
# Field code helpers
# ---------------------------------------------------------------------------

def _make_run(text, rpr_elem=None, preserve_space=False):
    """Create a <w:r> element with optional rPr and text."""
    r = etree.SubElement(etree.Element('dummy'), qn('w:r'))
    r.getparent().remove(r)
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    if preserve_space or (text and (text[0] == ' ' or text[-1] == ' ')):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _make_fld_char(fld_type, rpr_elem=None):
    """Create a <w:r> with <w:fldChar> (begin/separate/end)."""
    r = etree.SubElement(etree.Element('dummy'), qn('w:r'))
    r.getparent().remove(r)
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    fc = etree.SubElement(r, qn('w:fldChar'))
    fc.set(qn('w:fldCharType'), fld_type)
    return r


def _make_instr_text(instr, rpr_elem=None):
    """Create a <w:r> with <w:instrText>."""
    r = etree.SubElement(etree.Element('dummy'), qn('w:r'))
    r.getparent().remove(r)
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    it = etree.SubElement(r, qn('w:instrText'))
    it.text = instr
    it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _build_field(instr_text, display_text, rpr_elem=None):
    """Build a complete field code (begin + instrText + separate + display + end).
    Returns list of <w:r> elements to insert."""
    return [
        _make_fld_char('begin', rpr_elem),
        _make_instr_text(instr_text, rpr_elem),
        _make_fld_char('separate', rpr_elem),
        _make_run(display_text, rpr_elem),
        _make_fld_char('end', rpr_elem),
    ]


# ---------------------------------------------------------------------------
# CITATION field codes — replace in-text citations
# ---------------------------------------------------------------------------

def _safe_para_text(para):
    """Get the concatenated text of all runs, treating None as ''."""
    parts = []
    for r in para.runs:
        parts.append(r.text or '')
    return ''.join(parts)


def _find_citation_in_runs(para, cite_text_with_parens):
    """Find which run(s) contain the citation text. Returns (run_index, char_offset)
    for start and end, or None if not found."""
    full_text = _safe_para_text(para)
    pos = full_text.find(cite_text_with_parens)
    if pos == -1:
        return None

    runs = list(para.runs)
    char_pos = 0
    start_run = start_offset = end_run = end_offset = None

    target_end = pos + len(cite_text_with_parens)

    for i, run in enumerate(runs):
        run_text = run.text or ''
        run_start = char_pos
        run_end = char_pos + len(run_text)
        if start_run is None and run_end > pos:
            start_run = i
            start_offset = pos - run_start
        if end_run is None and run_end >= target_end:
            end_run = i
            end_offset = target_end - run_start
            break
        char_pos = run_end

    if start_run is not None and end_run is not None:
        return (start_run, start_offset, end_run, end_offset)
    return None


def _split_run_at(run_elem, offset):
    """Split a run element at the given character offset.
    Returns (left_run, right_run). The original run is replaced."""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    text_el = run_elem.find(qn('w:t'))
    if text_el is None:
        return run_elem, None
    full = text_el.text or ''
    left_text = full[:offset]
    right_text = full[offset:]

    rpr = run_elem.find(qn('w:rPr'))

    left_run = copy.deepcopy(run_elem)
    lt = left_run.find(qn('w:t'))
    lt.text = left_text
    if left_text and (left_text[0] == ' ' or left_text[-1] == ' '):
        lt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    right_run = copy.deepcopy(run_elem)
    rt = right_run.find(qn('w:t'))
    rt.text = right_text
    if right_text and (right_text[0] == ' ' or right_text[-1] == ' '):
        rt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return left_run, right_run


def fix_citations(doc, bib_entries, tex_path):
    """Replace in-text parenthetical citations with CITATION field codes."""
    citation_map = _build_citation_map(bib_entries, tex_path)

    rendered_to_keys = {}
    for rendered, keys in citation_map:
        paren_form = f'({rendered})'
        if paren_form not in rendered_to_keys:
            rendered_to_keys[paren_form] = keys

    replaced = 0
    for para in doc.paragraphs:
        if not para.runs:
            continue
        style_name = para.style.name if para.style else ''
        if style_name in ('Bibliography', 'Title', 'Author', 'Abstract Title'):
            continue

        full_text = _safe_para_text(para)
        replacements = []
        for cite_text, keys in rendered_to_keys.items():
            start = 0
            while True:
                pos = full_text.find(cite_text, start)
                if pos == -1:
                    break
                replacements.append((pos, pos + len(cite_text), cite_text, keys))
                start = pos + len(cite_text)

        if not replacements:
            continue

        replacements.sort(key=lambda x: x[0], reverse=True)

        for start_pos, end_pos, cite_text, keys in replacements:
            loc = _find_citation_in_runs(para, cite_text)
            if loc is None:
                continue

            start_run_idx, start_off, end_run_idx, end_off = loc
            runs = list(para.runs)
            p_elem = para._element

            rpr_elem = None
            first_run = runs[start_run_idx]._element
            existing_rpr = first_run.find(qn('w:rPr'))
            if existing_rpr is not None:
                rpr_elem = copy.deepcopy(existing_rpr)

            if len(keys) == 1:
                instr = f' CITATION {keys[0]} \\l 1033 '
            else:
                instr = f' CITATION {keys[0]} \\l 1033 '
                for k in keys[1:]:
                    instr += f'\\m {k} '

            field_runs = _build_field(instr, cite_text, rpr_elem)

            run_elems = [r._element for r in runs]

            if start_run_idx == end_run_idx:
                target_run = run_elems[start_run_idx]
                text_el = target_run.find(qn('w:t'))
                full_run_text = text_el.text or ''

                before_text = full_run_text[:start_off]
                after_text = full_run_text[end_off:]

                parent = target_run.getparent()
                insert_point = list(parent).index(target_run)
                parent.remove(target_run)

                elems_to_insert = []
                if before_text:
                    before_run = _make_run(before_text, rpr_elem, preserve_space=True)
                    elems_to_insert.append(before_run)
                elems_to_insert.extend(field_runs)
                if after_text:
                    after_run = _make_run(after_text, rpr_elem, preserve_space=True)
                    elems_to_insert.append(after_run)

                for j, elem in enumerate(elems_to_insert):
                    parent.insert(insert_point + j, elem)

            else:
                first_elem = run_elems[start_run_idx]
                last_elem = run_elems[end_run_idx]
                parent = first_elem.getparent()

                first_text_el = first_elem.find(qn('w:t'))
                before_text = (first_text_el.text or '')[:start_off]

                last_text_el = last_elem.find(qn('w:t'))
                after_text = (last_text_el.text or '')[end_off:]

                insert_point = list(parent).index(first_elem)
                for idx in range(start_run_idx, end_run_idx + 1):
                    parent.remove(run_elems[idx])

                elems_to_insert = []
                if before_text:
                    elems_to_insert.append(_make_run(before_text, rpr_elem, preserve_space=True))
                elems_to_insert.extend(field_runs)
                if after_text:
                    elems_to_insert.append(_make_run(after_text, rpr_elem, preserve_space=True))

                for j, elem in enumerate(elems_to_insert):
                    parent.insert(insert_point + j, elem)

            replaced += 1

    print(f"  Replaced {replaced} citation groups with CITATION field codes")


# ---------------------------------------------------------------------------
# BIBLIOGRAPHY field — replace reference list
# ---------------------------------------------------------------------------

def fix_bibliography_field(doc):
    """Replace all Bibliography-styled paragraphs with a single BIBLIOGRAPHY field."""
    body = doc.element.find(qn('w:body'))
    bib_paras = []
    refs_heading = None

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if style_name == 'Bibliography':
            bib_paras.append(para._element)
        elif style_name == 'Heading 1':
            clean = re.sub(r'^\d+(?:\.\d+)*\.?\s+', '',
                           _safe_para_text(para).strip().upper())
            if clean == 'REFERENCES':
                refs_heading = para._element

    if not bib_paras:
        print("  No Bibliography paragraphs found — skipping BIBLIOGRAPHY field")
        return

    first_bib = bib_paras[0]
    insert_idx = list(body).index(first_bib)

    # If no REFERENCES heading exists, create one before the bibliography
    if refs_heading is None:
        heading_p = etree.SubElement(etree.Element('dummy'), qn('w:p'))
        heading_p.getparent().remove(heading_p)
        h_ppr = etree.SubElement(heading_p, qn('w:pPr'))
        # Style as Heading 1
        h_pstyle = etree.SubElement(h_ppr, qn('w:pStyle'))
        h_pstyle.set(qn('w:val'), 'Heading1')
        # Page break before
        etree.SubElement(h_ppr, qn('w:pageBreakBefore'))
        set_spacing(h_ppr, line=LINE_SPACING_SINGLE,
                    before=HEADING_SPACING_BEFORE, after=HEADING_SPACING_AFTER)
        set_alignment(h_ppr, 'center')
        clear_first_indent(h_ppr)
        # Add run with "REFERENCES" text
        h_rpr = make_clean_rpr(BODY_FONT, 12, bold=True, italic=False, color=BLACK)
        h_run = _make_run('REFERENCES', h_rpr)
        heading_p.append(h_run)
        body.insert(insert_idx, heading_p)
        insert_idx += 1
        print("  Created REFERENCES heading with page break")

    # Build BIBLIOGRAPHY field paragraph
    rpr = make_clean_rpr(BODY_FONT, BODY_SIZE_PT, bold=False, italic=False, color=BLACK)
    field_elems = _build_field(' BIBLIOGRAPHY ', '[Update field to generate bibliography]', rpr)

    bib_p = etree.SubElement(etree.Element('dummy'), qn('w:p'))
    bib_p.getparent().remove(bib_p)
    ppr = etree.SubElement(bib_p, qn('w:pPr'))
    set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0, after=0)
    for r in field_elems:
        bib_p.append(r)

    body.insert(insert_idx, bib_p)

    for bp in bib_paras:
        body.remove(bp)

    print(f"  Replaced {len(bib_paras)} bibliography entries with BIBLIOGRAPHY field")


# ---------------------------------------------------------------------------
# SEQ fields for caption numbering + bookmarks
# ---------------------------------------------------------------------------

_next_bookmark_id = 100


def _get_bookmark_id():
    global _next_bookmark_id
    _next_bookmark_id += 1
    return _next_bookmark_id


def fix_caption_fields(doc):
    """Replace hard-coded 'Figure N.' and 'Table N.' in captions with SEQ fields
    and bookmarks for cross-referencing."""
    fig_bookmarks = {}
    tab_bookmarks = {}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''

        if style_name == 'Image Caption':
            m = re.match(r'^(Figure\s+)(\d+)(\.?\s*)', para.text)
            if not m:
                continue
            fig_num = int(m.group(2))
            _replace_caption_with_seq(para, 'Figure', fig_num, fig_bookmarks)

        elif style_name == 'Table Caption':
            m = re.match(r'^(Table\s+)(\d+)(\.?\s*)', para.text)
            if not m:
                continue
            tab_num = int(m.group(2))
            _replace_caption_with_seq(para, 'Table', tab_num, tab_bookmarks)

    print(f"  Added SEQ fields: {len(fig_bookmarks)} figures, {len(tab_bookmarks)} tables")
    return fig_bookmarks, tab_bookmarks


def _replace_caption_with_seq(para, seq_name, num, bookmark_dict):
    """Replace 'Figure N. ' or 'Table N. ' prefix with SEQ field + bookmark."""
    p_elem = para._element
    runs = list(para.runs)
    if not runs:
        return

    first_run = runs[0]._element
    text_el = first_run.find(qn('w:t'))
    if text_el is None:
        return

    full_text = text_el.text or ''
    pat = re.compile(rf'^({seq_name}\s+)(\d+)(\.?\s*)')
    m = pat.match(full_text)
    if not m:
        return

    prefix = m.group(1)
    suffix_dot = m.group(3)
    rest_text = full_text[m.end():]

    rpr_elem = first_run.find(qn('w:rPr'))
    rpr_copy = copy.deepcopy(rpr_elem) if rpr_elem is not None else None

    bm_id = _get_bookmark_id()
    bm_name = f'_Ref_{seq_name.lower()}{num}'
    bookmark_dict[num] = bm_name

    parent = first_run.getparent()
    insert_idx = list(parent).index(first_run)
    parent.remove(first_run)

    elems = []

    bm_start = etree.SubElement(etree.Element('dummy'), qn('w:bookmarkStart'))
    bm_start.getparent().remove(bm_start)
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), bm_name)
    elems.append(bm_start)

    elems.append(_make_run(prefix, rpr_copy, preserve_space=True))

    seq_instr = f' SEQ {seq_name} \\* ARABIC '
    elems.extend(_build_field(seq_instr, str(num), rpr_copy))

    bm_end = etree.SubElement(etree.Element('dummy'), qn('w:bookmarkEnd'))
    bm_end.getparent().remove(bm_end)
    bm_end.set(qn('w:id'), str(bm_id))
    elems.append(bm_end)

    if suffix_dot:
        elems.append(_make_run(suffix_dot, rpr_copy, preserve_space=True))

    if rest_text:
        elems.append(_make_run(rest_text, rpr_copy, preserve_space=True))

    for j, elem in enumerate(elems):
        parent.insert(insert_idx + j, elem)


# ---------------------------------------------------------------------------
# REF cross-references — replace "Figure 1" / "Table 1" in body text
# ---------------------------------------------------------------------------

def fix_cross_references(doc, fig_bookmarks, tab_bookmarks):
    """Replace cross-references in body text with REF fields.

    Handles two patterns:
    1. 'Figure N' / 'Table N' — when pandoc-crossref was used (has numbers)
    2. 'Figure ' / 'Table ' followed by nothing — when pandoc dropped \\ref{}
       (just the label word with trailing space and no number)

    For pattern 2, assigns numbers sequentially (1st Figure ref → Figure 1, etc.).
    """
    ref_count = 0
    # Pattern 1: explicit numbers
    numbered_pattern = re.compile(r'(Figure|Table)\s+(\d+)')
    # Pattern 2: bare labels from failed \ref{} — pandoc renders \ref{} as empty,
    # leaving "Figure\xa0" possibly followed by space.  The \xa0 may be the ONLY
    # trailing whitespace (e.g. "Table\xa0." where the period follows immediately).
    # Match: label word + \xa0 + optional extra whitespace.
    bare_pattern = re.compile('(Figure|Table) [ ]*')

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if style_name in ('Image Caption', 'Table Caption', 'Captioned Figure',
                          'Bibliography', 'Title', 'Author'):
            continue

        full_text = _safe_para_text(para)

        # Try numbered pattern first
        matches = list(numbered_pattern.finditer(full_text))
        for match in reversed(matches):
            ref_type = match.group(1)
            ref_num = int(match.group(2))
            ref_text = match.group(0)

            if ref_type == 'Figure':
                bm_name = fig_bookmarks.get(ref_num)
            else:
                bm_name = tab_bookmarks.get(ref_num)

            if not bm_name:
                continue

            if _replace_ref_in_para(para, ref_text, bm_name, ref_text):
                ref_count += 1

        if matches:
            continue

        # Try bare pattern (pandoc dropped \ref{})
        bare_matches = list(bare_pattern.finditer(full_text))
        # Process forward (not reversed) because _find_citation_in_runs uses
        # str.find() which returns the FIRST occurrence.  Forward order ensures
        # each replacement removes the leftmost match so the next find() hits
        # the next one.
        for match in bare_matches:
            ref_type = match.group(1)
            ref_text = match.group(0)  # e.g. "Figure\xa0 " or "Table\xa0"

            # When only one bookmark of a type exists, all bare refs map to it.
            # Otherwise assign sequentially within the paragraph.
            if ref_type == 'Figure':
                if len(fig_bookmarks) == 1:
                    ref_num = next(iter(fig_bookmarks))
                else:
                    ref_num = sum(1 for m in bare_matches
                                  if m.group(1) == 'Figure' and m.start() <= match.start())
                bm_name = fig_bookmarks.get(ref_num)
                display = f"Figure {ref_num}"
            else:
                if len(tab_bookmarks) == 1:
                    ref_num = next(iter(tab_bookmarks))
                else:
                    ref_num = sum(1 for m in bare_matches
                                  if m.group(1) == 'Table' and m.start() <= match.start())
                bm_name = tab_bookmarks.get(ref_num)
                display = f"Table {ref_num}"

            if not bm_name:
                continue

            if _replace_ref_in_para(para, ref_text, bm_name, display):
                ref_count += 1

    print(f"  Replaced {ref_count} cross-references with REF field codes")


def _replace_ref_in_para(para, search_text, bm_name, display_text):
    """Replace search_text in paragraph with a REF field pointing to bm_name."""
    loc = _find_citation_in_runs(para, search_text)
    if loc is None:
        return False

    start_run_idx, start_off, end_run_idx, end_off = loc
    runs = list(para.runs)
    run_elems = [r._element for r in runs]

    rpr_elem = None
    first_run_el = run_elems[start_run_idx]
    existing_rpr = first_run_el.find(qn('w:rPr'))
    if existing_rpr is not None:
        rpr_elem = copy.deepcopy(existing_rpr)

    # \* MERGEFORMAT makes the REF result inherit LOCAL formatting (non-bold
    # body text) rather than the formatting from the bookmark target (which
    # is inside a bold caption).
    instr = f' REF {bm_name} \\h \\* MERGEFORMAT '
    # Strip bold from rpr so cross-refs render plain (not bold like captions)
    if rpr_elem is not None:
        for bold_tag in (qn('w:b'), qn('w:bCs')):
            b = rpr_elem.find(bold_tag)
            if b is not None:
                rpr_elem.remove(b)
        # Explicitly set w:b val="false" to override any inherited bold
        b_off = etree.SubElement(rpr_elem, qn('w:b'))
        b_off.set(qn('w:val'), '0')
        bcs_off = etree.SubElement(rpr_elem, qn('w:bCs'))
        bcs_off.set(qn('w:val'), '0')
    field_runs = _build_field(instr, display_text, rpr_elem)

    # Helper: determine the character immediately following the replaced text.
    # If it's a word character and the replacement consumed all whitespace,
    # we need to insert a space between the field and the next word.
    def _needs_trailing_space(after_txt, end_ri, end_o):
        """Check if we need a space after the REF field."""
        if after_txt and after_txt[0].isalnum():
            return True
        if not after_txt and end_ri + 1 < len(run_elems):
            next_el = run_elems[end_ri + 1]
            next_t = next_el.find(qn('w:t'))
            if next_t is not None and next_t.text and next_t.text[0].isalnum():
                return True
        return False

    if start_run_idx == end_run_idx:
        target_run = run_elems[start_run_idx]
        text_el = target_run.find(qn('w:t'))
        full_run_text = text_el.text or ''

        before_text = full_run_text[:start_off]
        after_text = full_run_text[end_off:]

        need_space = _needs_trailing_space(after_text, end_run_idx, end_off)

        parent = target_run.getparent()
        insert_point = list(parent).index(target_run)
        parent.remove(target_run)

        elems_to_insert = []
        if before_text:
            elems_to_insert.append(_make_run(before_text, rpr_elem, preserve_space=True))
        elems_to_insert.extend(field_runs)
        if need_space:
            elems_to_insert.append(_make_run(' ', rpr_elem, preserve_space=True))
        if after_text:
            elems_to_insert.append(_make_run(after_text, rpr_elem, preserve_space=True))

        for j, elem in enumerate(elems_to_insert):
            parent.insert(insert_point + j, elem)

        return True

    else:
        first_elem = run_elems[start_run_idx]
        last_elem = run_elems[end_run_idx]
        parent = first_elem.getparent()

        first_text_el = first_elem.find(qn('w:t'))
        before_text = (first_text_el.text or '')[:start_off]

        last_text_el = last_elem.find(qn('w:t'))
        after_text = (last_text_el.text or '')[end_off:]

        need_space = _needs_trailing_space(after_text, end_run_idx, end_off)

        insert_point = list(parent).index(first_elem)
        for idx in range(start_run_idx, end_run_idx + 1):
            parent.remove(run_elems[idx])

        elems_to_insert = []
        if before_text:
            elems_to_insert.append(_make_run(before_text, rpr_elem, preserve_space=True))
        elems_to_insert.extend(field_runs)
        if need_space:
            elems_to_insert.append(_make_run(' ', rpr_elem, preserve_space=True))
        if after_text:
            elems_to_insert.append(_make_run(after_text, rpr_elem, preserve_space=True))

        for j, elem in enumerate(elems_to_insert):
            parent.insert(insert_point + j, elem)

        return True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def apply_preprint_overrides(doc):
    """Convert a manuscript-formatted doc to preprint style in-place.

    Changes from manuscript defaults:
    - Body text: single-spaced, 11pt, 12pt paragraph spacing instead of indent
    - Margins: 0.75 inch on all sides
    - Headings and tables: unchanged
    """
    PREPRINT_SIZE = 11
    PREPRINT_SPACING_AFTER = 160  # ~8pt spacing between paragraphs

    # Override Normal style
    for s in doc.styles:
        if s.name == "Normal":
            ppr = s.element.find(qn("w:pPr"))
            if ppr is None:
                ppr = etree.SubElement(s.element, qn("w:pPr"))
            set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0,
                        after=PREPRINT_SPACING_AFTER)
            # Update font size in style rPr
            rpr = s.element.find(qn("w:rPr"))
            if rpr is not None:
                for tag in (qn("w:sz"), qn("w:szCs")):
                    el = rpr.find(tag)
                    if el is not None:
                        el.set(qn("w:val"), str(PREPRINT_SIZE * 2))
            break

    # Override every body paragraph
    for para in doc.paragraphs:
        if is_in_table(para):
            continue
        style_name = para.style.name if para.style else "Normal"
        if style_name.startswith("Heading") or style_name in (
            "Title", "Author", "Abstract Title", "Abstract", "Bibliography",
            "Image Caption", "Table Caption", "Captioned Figure",
        ):
            continue
        ppr = ensure_ppr(para)
        set_spacing(ppr, line=LINE_SPACING_SINGLE, before=0,
                    after=PREPRINT_SPACING_AFTER)
        clear_first_indent(ppr)
        apply_font_to_runs(para, font=BODY_FONT, size_pt=PREPRINT_SIZE)

    # Tighter margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    print("  Applied preprint overrides (single-spaced, 11pt, 0.75in margins)")


def format_docx(input_path, output_path=None, hires_image=None,
                bib_path=None, tex_path=None, anonymize=False,
                style="manuscript"):
    if output_path is None:
        output_path = input_path

    print(f"Processing: {input_path} (style={style})")
    doc = Document(input_path)

    fix_theme(doc)
    fix_normal_style(doc)
    fix_styles(doc)
    fix_page_layout(doc)
    fix_heading_numbering(doc)
    fix_keywords_and_pagebreak(doc, tex_path=tex_path)
    fix_paragraphs(doc, anonymize=anonymize)

    fix_captions(doc)                        # 1. Add "Figure N." / "Table N." prefixes
    fig_bm, tab_bm = fix_caption_fields(doc) # 2. Replace prefixes with SEQ fields
    fix_cross_references(doc, fig_bm, tab_bm)# 3. Replace body "Figure 1" with REF fields
    fix_tables(doc)

    if bib_path and tex_path and Path(bib_path).exists() and Path(tex_path).exists():
        bib_entries = parse_bib(bib_path)
        fix_citations(doc, bib_entries, tex_path)
        fix_bibliography_field(doc)

    if hires_image and Path(hires_image).exists():
        replace_image(doc, hires_image)

    if style == "preprint":
        apply_preprint_overrides(doc)

    doc.save(output_path)

    if bib_path and tex_path and Path(bib_path).exists() and Path(tex_path).exists():
        _inject_bibliography_xml(output_path, bib_entries)

    print(f"Saved: {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("input", help="Input DOCX file")
    parser.add_argument("output", nargs="?", default=None,
                        help="Output DOCX file (default: overwrite input)")
    parser.add_argument("--image", default=None,
                        help="High-res image to replace the first embedded image")
    parser.add_argument("--bib", default=None,
                        help="BibTeX file for Word Reference Manager integration")
    parser.add_argument("--tex", default=None,
                        help="LaTeX source file for citation key mapping")
    parser.add_argument("--anonymize", action="store_true",
                        help="Remove the author block (double-blind submission)")
    args = parser.parse_args()
    format_docx(args.input, args.output, hires_image=args.image,
                bib_path=args.bib, tex_path=args.tex, anonymize=args.anonymize)
