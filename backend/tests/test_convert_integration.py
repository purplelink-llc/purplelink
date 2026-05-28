import shutil
from pathlib import Path

import pytest

from latextools import runner

FIXTURES = Path(__file__).parent / "fixtures"
has_pandoc = pytest.mark.skipif(
    shutil.which("pandoc") is None, reason="pandoc not installed (run inside image)"
)


@has_pandoc
def test_convert_returns_docx(tmp_path):
    tex = (FIXTURES / "manuscript.tex").read_text()
    result = runner.convert_to_manuscript(tmp_path, tex, anonymize=False)
    assert result.ok is True
    # .docx is a zip; magic bytes "PK"
    assert result.docx_bytes[:2] == b"PK"


@has_pandoc
def test_convert_anonymize_removes_author(tmp_path):
    import io
    from docx import Document

    tex = (FIXTURES / "manuscript.tex").read_text()
    anon = runner.convert_to_manuscript(tmp_path, tex, anonymize=True)
    named = runner.convert_to_manuscript(tmp_path, tex, anonymize=False)
    anon_text = "\n".join(p.text for p in Document(io.BytesIO(anon.docx_bytes)).paragraphs)
    named_text = "\n".join(p.text for p in Document(io.BytesIO(named.docx_bytes)).paragraphs)
    assert "Jane Researcher" in named_text
    assert "Jane Researcher" not in anon_text
